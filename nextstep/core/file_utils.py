"""
File utilities for parsing resume documents.

Supports:
- PDF files (using PyPDF2)
- DOCX files (using python-docx)
"""

import logging
from typing import Optional
from io import BytesIO

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        file: File-like object or uploaded file
        
    Returns:
        Extracted text as string
    """
    try:
        from PyPDF2 import PdfReader
        
        # Handle Django UploadedFile
        if hasattr(file, 'read'):
            file_content = file.read()
            file.seek(0)  # Reset for potential re-reads
            pdf_file = BytesIO(file_content)
        else:
            pdf_file = file
            
        reader = PdfReader(pdf_file)
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise ValueError(f"Could not read PDF file: {str(e)}")


def extract_text_from_docx(file) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        file: File-like object or uploaded file
        
    Returns:
        Extracted text as string
    """
    try:
        from docx import Document
        
        # Handle Django UploadedFile
        if hasattr(file, 'read'):
            file_content = file.read()
            file.seek(0)  # Reset for potential re-reads
            docx_file = BytesIO(file_content)
        else:
            docx_file = file
            
        doc = Document(docx_file)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Could not read DOCX file: {str(e)}")


def extract_resume_text(file, content_type: Optional[str] = None, filename: Optional[str] = None) -> str:
    """
    Extract text from a resume file based on its type.
    
    Args:
        file: Uploaded file object
        content_type: MIME type of the file
        filename: Original filename (used as fallback for type detection)
        
    Returns:
        Extracted text as string
        
    Raises:
        ValueError: If file type is not supported
    """
    # Determine file type from content_type or filename
    if content_type:
        content_type = content_type.lower()
        
        if 'pdf' in content_type:
            return extract_text_from_pdf(file)
        elif 'wordprocessingml' in content_type or 'docx' in content_type:
            return extract_text_from_docx(file)
    
    # Fallback to filename extension
    if filename:
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return extract_text_from_pdf(file)
        elif filename_lower.endswith('.docx'):
            return extract_text_from_docx(file)
        elif filename_lower.endswith('.doc'):
            raise ValueError("Legacy .doc format is not supported. Please use .docx format.")
    
    raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")
